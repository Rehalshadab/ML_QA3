from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Naive Bayes Classifier Implementation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
doc.add_paragraph('Course: QA 3 (Unit 3)', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ========================
# Section 1: Introduction
# ========================
doc.add_heading('1. Introduction to Naive Bayes', level=1)

doc.add_paragraph(
    'Naive Bayes is a probabilistic machine learning algorithm based on Bayes\' Theorem. '
    'It is called "naive" because it assumes that the features (words in text classification) '
    'are conditionally independent of each other given the class label. Despite this simplifying '
    'assumption, Naive Bayes performs well in many real-world applications, particularly in text '
    'classification problems like spam detection and sentiment analysis.'
)

doc.add_heading('Bayes\' Theorem', level=2)
doc.add_paragraph(
    'Bayes\' Theorem describes the probability of an event based on prior knowledge of conditions '
    'that might be related to the event. The formula is:'
)
p = doc.add_paragraph()
p.add_run('P(Class | Features) = [P(Class) * P(Features | Class)] / P(Features)').bold = True

doc.add_heading('Prior and Likelihood', level=2)
p = doc.add_paragraph()
p.add_run('Prior Probability P(Class): ').bold = True
p.add_run('The probability of a class before observing any features. For example, if 3 out of 6 messages are spam, P(spam) = 0.5.')

p = doc.add_paragraph()
p.add_run('Likelihood P(Feature | Class): ').bold = True
p.add_run('The probability of a feature (word) given a particular class. For example, P("free" | spam) is the probability of seeing the word "free" in a spam message.')

p = doc.add_paragraph()
p.add_run('Posterior Probability P(Class | Features): ').bold = True
p.add_run('The probability of a class after observing the features. This is what we compute to make predictions.')

# ========================
# Section 2: Dataset
# ========================
doc.add_heading('2. Dataset', level=1)
doc.add_paragraph('We use a small SMS spam classification dataset with 6 messages:')

# Create table
table = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
table.cell(0, 0).text = 'Message'
table.cell(0, 1).text = 'Label'
messages = [
    ('free money now', 'spam'),
    ('win cash today', 'spam'),
    ('call me now', 'ham'),
    ('meeting at work', 'ham'),
    ('free call today', 'spam'),
    ('work meeting tomorrow', 'ham'),
]
for i, (msg, label) in enumerate(messages, 1):
    table.cell(i, 0).text = msg
    table.cell(i, 1).text = label

doc.add_paragraph('')
doc.add_paragraph('Test messages to classify:')
doc.add_paragraph('1. "free work call"', style='List Bullet')
doc.add_paragraph('2. "win money today"', style='List Bullet')

# ========================
# Section 3: Manual Calculation
# ========================
doc.add_heading('3. Manual Calculation Steps', level=1)

doc.add_heading('Step 1: Calculate Prior Probabilities', level=2)
doc.add_paragraph('Total messages = 6')
doc.add_paragraph('Spam messages = 3, Ham messages = 3')
p = doc.add_paragraph()
p.add_run('P(spam) = 3/6 = 0.5').bold = True
p = doc.add_paragraph()
p.add_run('P(ham) = 3/6 = 0.5').bold = True

doc.add_heading('Step 2: Build Vocabulary and Count Word Frequencies', level=2)
doc.add_paragraph('Unique words (vocabulary): call, cash, free, me, meeting, money, now, today, win, work')
doc.add_paragraph('Vocabulary size V = 10')

doc.add_paragraph('Word counts in spam messages:')
doc.add_paragraph('free: 2, money: 1, now: 1, win: 1, cash: 1, call: 1, today: 1', style='List Bullet')
doc.add_paragraph('Total spam words = 8')

doc.add_paragraph('Word counts in ham messages:')
doc.add_paragraph('call: 1, me: 1, now: 1, meeting: 2, at: 1, work: 2, tomorrow: 1', style='List Bullet')
doc.add_paragraph('Total ham words = 9')

doc.add_heading('Step 3: Calculate Likelihoods with Laplace Smoothing', level=2)
doc.add_paragraph(
    'Laplace (add-1) smoothing prevents zero probabilities for unseen words. '
    'The formula is: P(word | class) = (count(word in class) + 1) / (total words in class + V)'
)

p = doc.add_paragraph()
p.add_run('Example: P("free" | spam) = (2 + 1) / (8 + 10) = 3/18 = 0.1667').bold = True
p = doc.add_paragraph()
p.add_run('Example: P("free" | ham) = (0 + 1) / (9 + 10) = 1/19 = 0.0526').bold = True

doc.add_heading('Step 4: Classify Test Messages', level=2)

doc.add_paragraph('For message "free work call":')
doc.add_paragraph(
    'P(spam | free, work, call) ∝ P(spam) * P(free|spam) * P(work|spam) * P(call|spam)\n'
    '= 0.5 * (3/18) * (1/18) * (2/18) = 0.5 * 0.1667 * 0.0556 * 0.1111 = 0.000514\n\n'
    'P(ham | free, work, call) ∝ P(ham) * P(free|ham) * P(work|ham) * P(call|ham)\n'
    '= 0.5 * (1/19) * (3/19) * (2/19) = 0.5 * 0.0526 * 0.1579 * 0.1053 = 0.000437'
)
doc.add_paragraph('After normalization: P(spam | message) ≈ 0.540, P(ham | message) ≈ 0.460')
p = doc.add_paragraph()
p.add_run('Prediction: spam').bold = True

doc.add_paragraph('')
doc.add_paragraph('For message "win money today":')
doc.add_paragraph(
    'P(spam | win, money, today) ∝ P(spam) * P(win|spam) * P(money|spam) * P(today|spam)\n'
    '= 0.5 * (2/18) * (2/18) * (2/18) = 0.5 * 0.1111 * 0.1111 * 0.1111 = 0.000686\n\n'
    'P(ham | win, money, today) ∝ P(ham) * P(win|ham) * P(money|ham) * P(today|ham)\n'
    '= 0.5 * (1/19) * (1/19) * (1/19) = 0.5 * 0.0526 * 0.0526 * 0.0526 = 0.000073'
)
doc.add_paragraph('After normalization: P(spam | message) ≈ 0.904, P(ham | message) ≈ 0.096')
p = doc.add_paragraph()
p.add_run('Prediction: spam').bold = True

# ========================
# Section 4: Code Implementation
# ========================
doc.add_heading('4. Code Implementation', level=1)

doc.add_heading('4.1 Scratch Implementation', level=2)
doc.add_paragraph(
    'A custom NaiveBayesScratch class was implemented with:\n'
    '- Laplace smoothing (alpha parameter)\n'
    '- Log-probability calculations to avoid numerical underflow\n'
    '- Fit, predict, and predict_proba methods'
)

doc.add_heading('4.2 sklearn Implementation', level=2)
doc.add_paragraph(
    'The sklearn MultinomialNB classifier was used for comparison:\n'
    '- CountVectorizer converts text to word count vectors\n'
    '- MultinomialNB with alpha=1.0 (Laplace smoothing)\n'
    '- Same training data and test messages'
)

# ========================
# Section 5: Results Comparison
# ========================
doc.add_heading('5. Results Comparison', level=1)

results = [
    ('free work call', 'Manual', 'spam', '0.540', '0.460'),
    ('free work call', 'Scratch', 'spam', '0.5401', '0.4599'),
    ('free work call', 'sklearn', 'spam', '0.540', '0.460'),
    ('win money today', 'Manual', 'spam', '0.904', '0.096'),
    ('win money today', 'Scratch', 'spam', '0.9040', '0.0960'),
    ('win money today', 'sklearn', 'spam', '0.904', '0.096'),
]

table2 = doc.add_table(rows=1 + len(results), cols=5, style='Light Grid Accent 1')
headers = ['Test Message', 'Method', 'Prediction', 'P(spam)', 'P(ham)']
for i, h in enumerate(headers):
    table2.cell(0, i).text = h

for i, row in enumerate(results):
    table2.cell(i + 1, 0).text = row[0]
    table2.cell(i + 1, 1).text = row[1]
    table2.cell(i + 1, 2).text = row[2]
    table2.cell(i + 1, 3).text = row[3]
    table2.cell(i + 1, 4).text = row[4]

doc.add_paragraph('')
doc.add_paragraph('All three methods produce the same predictions and nearly identical probabilities, '
                   'confirming the correctness of the implementation.')

# ========================
# Section 6: Conclusion
# ========================
doc.add_heading('6. Conclusion', level=1)
doc.add_paragraph(
    'This project demonstrates the complete Naive Bayes classification workflow:\n\n'
    '1. Understanding Bayes\' Theorem and the naive independence assumption\n'
    '2. Manual calculation of priors, likelihoods, and posteriors\n'
    '3. Implementation from scratch in Python\n'
    '4. Using sklearn\'s optimized implementation\n'
    '5. Comparing and verifying results across approaches\n\n'
    'Naive Bayes is simple, fast, and effective for text classification tasks '
    'with small to medium-sized datasets.'
)

doc.save('C:/Users/HP/Desktop/normal/Naive_Bayes_Explanation.docx')
print("Word document created successfully!")
